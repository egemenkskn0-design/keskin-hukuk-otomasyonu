import streamlit as st
import sqlite3
import hashlib
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime

# --- Şifre ve Word Fonksiyonları ---
def hash_sifre(sifre): return hashlib.sha256(sifre.encode()).hexdigest()

def word_uret(v):
    doc = docx.Document()
    p = doc.add_paragraph("KİRA SÖZLEŞMESİ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    doc.add_paragraph(f"Tarih: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"KİRAYA VEREN: {v['kiralayan']} (TC: {v['kiralayan_tc']})\nAdres: {v['kiralayan_adres']}\nIBAN: {v['kiralayan_iban']}")
    doc.add_paragraph(f"KİRACI: {v['kiraci']} (TC: {v['kiraci_tc']})\nAdres: {v['kiraci_adres']}")
    doc.add_paragraph(f"ŞARTLAR: {v['bedel']} {v['para_birimi']} - Başlangıç: {v['baslangic']} - Ödeme Günü: {v['odeme_gunu']}")
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

st.set_page_config(page_title="AV. EGEMEN KESKİN - YÖNETİM", layout="wide")
st.title("🔒 AV. EGEMEN KESKİN - YÖNETİCİ PANELİ")

# Şifre Kontrolü
conn = sqlite3.connect("hukuk_otomasyon.db")
c = conn.cursor()
c.execute("SELECT password_hash FROM admin_auth")
auth = c.fetchone()
conn.close()

giris = st.text_input("Yönetici Şifresi:", type="password")

if auth and hash_sifre(giris) == auth[0]:
    st.success("Erişim Onaylandı")
    tab1, tab2 = st.tabs(["📬 Yeni Sözleşme Talepleri", "📂 Mevcut Sözleşme İncelemeleri"])
    
    with tab1:
        st.subheader("Yeni Müvekkil Talepleri")
        conn = sqlite3.connect("hukuk_otomasyon.db")
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        c.execute("SELECT * FROM yeni_kontratlar ORDER BY id DESC")
        kayitlar = c.fetchall()
        for k in kayitlar:
            with st.expander(f"Müvekkil: {k['kiraci']} - {k['tarih']}"):
                st.write(f"Kiraya Veren: {k['kiralayan']} | TC: {k['kiralayan_tc']}")
                st.write(f"Bedel: {k['bedel']} {k['para_birimi']}")
                data = word_uret(k)
                st.download_button("📜 Sözleşmeyi Word Olarak İndir", data, f"Sozlesme_{k['kiraci']}.docx")
        conn.close()

    with tab2:
        st.subheader("Mevcut Sözleşme İnceleme Listesi")
        conn = sqlite3.connect("hukuk_otomasyon.db")
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        c.execute("SELECT * FROM eski_kontratlar ORDER BY id DESC")
        eski = c.fetchall()
        for e in eski:
            st.info(f"**Müvekkil:** {e['kiraci']} | **Talep:** {e['notlar']} | **Dosya:** {e['dosya_adi']}")
        conn.close()
elif giris:
    st.error("Hatalı Şifre!")
