import streamlit as st
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime
import sqlite3
import hashlib

# =====================================================================
# GÜVENLİ VERİ TABANI ALTYAPISI
# =====================================================================
def hash_sifre(sifre):
    return hashlib.sha256(sifre.encode()).hexdigest()

def veri_tabani_hazirla():
    conn = sqlite3.connect("hukuk_otomasyon.db")
    cursor = conn.cursor()
    # Şifre Yönetim Tablosu
    cursor.execute("CREATE TABLE IF NOT EXISTS admin_auth (id INTEGER PRIMARY KEY, password_hash TEXT)")
    # Sıfırdan Üretilen Kontratlar Tablosu
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS yeni_kontratlar (
            id INTEGER PRIMARY KEY AUTOINCREMENT, tarih TEXT, kiralayan TEXT, kiralayan_tc TEXT, 
            kiralayan_adres TEXT, kiralayan_iban TEXT, kiraci TEXT, kiraci_tc TEXT, kiraci_adres TEXT, 
            cins TEXT, bedel REAL, para_birimi TEXT, depozito INTEGER, artis INTEGER, baslangic TEXT, odeme_gunu INTEGER
        )
    """)
    # Mevcut/Eski Kontratlar Arşivi Tablosu
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS eski_kontratlar (
            id INTEGER PRIMARY KEY AUTOINCREMENT, islem_tarihi TEXT, kiralayan TEXT, kiraci TEXT, 
            baslangic_tarihi TEXT, aylik_bedel REAL, para_birimi TEXT, dosya_adi TEXT, notlar TEXT
        )
    """)
    conn.commit()
    conn.close()

veri_tabani_hazirla()

# Sayfa ve Kurumsal Arayüz Ayarları
st.set_page_config(page_title="Keskin Hukuk Bürosu", page_icon="⚖️", layout="wide")

# =====================================================================
# SOL TARAFTAKİ SEÇİM ALANI (SIDEBAR)
# =====================================================================
st.sidebar.markdown("### 🏛️ Keskin Hukuk Portalı")
islem_tipi = st.sidebar.radio(
    "Lütfen Durumunuza Uygun Seçeneği İşaretleyin:",
    ["✨ Sıfırdan Kira Sözleşmesi Hazırlatmak İstiyorum", 
     "📂 Elimde İmzalanmış Mevcut Bir Sözleşme Var",
     "🔒 Avukat Yönetim Paneli (Sadece Ofis İçi Giriş)"]
)

# =====================================================================
# SENARYO 1: SIFIRDAN KİRA SÖZLEŞMESİ HAZIRLAMA FORMU
# =====================================================================
if islem_tipi == "✨ Sıfırdan Kira Sözleşmesi Hazırlatmak İstiyorum":
    st.title("⚖️ Keskin Hukuk Bürosu - Dijital Bilgi Portalı")
    st.markdown("##### Müvekkil / Sözleşme Tarafları Bilgi Giriş Ekranı")
    st.divider()
    st.subheader("📋 Sıfırdan Sözleşme Üretim Formu")
    st.info("Yeni bir kira sözleşmesi oluşturulması için lütfen aşağıdaki alanları resmi evraklarınıza uygun şekilde doldurunuz.")
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("##### **Kiraya Veren (Mal Sahibi) Bilgileri**")
        kiralayan_ad = st.text_input("Adı Soyadı / Unvanı", key="y_k_ad")
        kiralayan_tc = st.text_input("T.C. Kimlik No / Vergi No", key="y_k_tc")
        kiralayan_adres = st.text_area("Resmi Tebligat Adresi", key="y_k_adr")
        kiralayan_iban = st.text_input("Kira Ödemelerinin Yapılacağı IBAN Numarası", key="y_k_iban")

    with col2:
        st.markdown("##### **Kiracı Bilgileri**")
        kiraci_ad = st.text_input("Adı Soyadı / Şirket Unvanı", key="y_kr_ad")
        kiraci_tc = st.text_input("T.C. Kimlik / Pasaport / Vergi No", key="y_kr_tc")
        kiraci_adres = st.text_area("Kiracının İkametgah Adresi", key="y_kr_adr")

    st.divider()
    st.markdown("##### **Gayrimenkul ve Finansal Detaylar**")
    col3, col4 = st.columns(2)
    with col3:
        gayrimenkul_turu = st.selectbox("Kiralanan Taşınmazın Cinsi", ["Konut", "Çatılı İşyeri", "Arsa/Arazi"])
        
        sub_col1, sub_col2 = st.columns([3, 1])
        with sub_col1:
            kira_bedeli_aylik = st.number_input("Aylık Kira Bedeli", min_value=0, value=20000, step=500, key="y_bedel")
        with sub_col2:
            para_birimi = st.selectbox("Para Birimi", ["TL (₺)", "USD ($)", "EUR (€)"], key="y_pb")
            
        depozito_ay_sayisi = st.slider("Güvence Bedeli (Kaç Aylık Kira Karşılığı?)", min_value=0, max_value=12, value=2)
        # Güvence bedeli girişi (Mevcut slider'ın)
        depozito_ay_sayisi = st.slider("Güvence Bedeli (Kaç Aylık Kira Karşılığı?)", min_value=0, max_value=12, value=2)

        # HUKUKİ KONTROL (Buraya ekle!)
        if gayrimenkul_turu == "Çatılı İşyeri" and depozito_ay_sayisi > 3:
            st.error("⚠️ Türk Borçlar Kanunu md. 342 uyarınca, konut ve çatılı iş yeri kiralarında güvence bedeli 3 aylık kira bedelini aşamaz.")
        elif gayrimenkul_turu == "Konut" and depozito_ay_sayisi > 3:
            st.error("⚠️ Türk Borçlar Kanunu md. 342 uyarınca, konutlarda güvence bedeli 3 aylık kira bedelini aşamaz.")
    
    with col4:
        sozlesme_artis_orani = st.slider("Yeni Dönem Yıllık Artış Oranı (%)", min_value=0, max_value=150, value=50)
        kira_baslangic = st.date_input("Kira Sözleşmesi Başlangıç Tarihi", datetime.now(), key="y_tarih")
        odeme_gunu = st.number_input("Her Ayın Hangi Günü Ödenecek? (1-30)", min_value=1, max_value=30, value=5)

    if st.button("🚀 BİLGİLERİ AVUKATA GÖNDER", use_container_width=True):
        if not kiralayan_ad or not kiraci_ad:
            st.error("❌ Hata: İşlemin tamamlanabilmesi için Kiraya Veren ve Kiracı isim alanları boş bırakılamaz.")
        else:
            conn = sqlite3.connect("hukuk_otomasyon.db")
            cursor = conn.cursor()
            cursor.execute("""
    INSERT INTO yeni_kontratlar 
    (tarih, kiralayan, kiralayan_tc, kiralayan_adres, kiralayan_iban, kiraci, kiraci_tc, kiraci_adres, cins, bedel, para_birimi, depozito, artis, baslangic, odeme_gunu)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
""", (tarih_degeri, kiralayan_degeri, ...)) # Buraya da 15 tane değişken gelmeli
            """, (
                datetime.now().strftime("%d.%m.%Y %H:%M"), kiralayan_ad, kiralayan_tc, kiralayan_adres, kiralayan_iban,
                kiraci_ad, kiraci_tc, kiraci_adres, gayrimenkul_turu, kira_bedeli_aylik, para_birimi, depozito_ay_sayisi,
                sozlesme_artis_orani, kira_baslangic.strftime('%d.%m.%Y'), odeme_gunu
            ))
            conn.commit()
            conn.close()
            st.success("🎯 Bilgileriniz Avukat Egemen Keskin'in sistemine başarıyla aktarıldı. Teşekkür ederiz.")

# =====================================================================
# SENARYO 2: ELDEKİ MEVCUT SÖZLEŞMEYİ YÜKLEME FORMU
# =====================================================================
elif islem_tipi == "📂 Elimde İmzalanmış Mevcut Bir Sözleşme Var":
    st.title("⚖️ Keskin Hukuk Bürosu - Dijital Bilgi Portalı")
    st.markdown("##### Müvekkil / Sözleşme Tarafları Bilgi Giriş Ekranı")
    st.divider()
    st.subheader("📂 Mevcut Sözleşme Bildirim ve Yükleme Formu")
    st.info("DHA önceden imzalanmış sözleşmenizin hukuki analizi için lütfen alanları doldurup evrakı yükleyiniz.")
    
    col_e1, col_e2 = st.columns(2)
    with col_e1:
        e_kiralayan = st.text_input("Sözleşmedeki Kiraya Veren (Mal Sahibi):", key="e_k")
        e_kiraci = st.text_input("Sözleşmedeki Kiracı:", key="e_kr")
        e_baslangic = st.date_input("Sözleşmenin İmza / Başlangıç Tarihi:", datetime.now(), key="e_tar")
        
        e_sub1, e_sub2 = st.columns([3, 1])
        with e_sub1:
            e_bedel = st.number_input("Sözleşmede Yazılan Aylık Kira Bedeli:", min_value=0, value=15000, key="e_bedel")
        with e_sub2:
            e_pb = st.selectbox("Para Birimi", ["TL (₺)", "USD ($)", "EUR (€)"], key="e_pb")
            
    with col_e2:
        st.markdown("##### **Sözleşme Görseli / PDF Yükleme**")
        yuklenen_dosya = st.file_uploader("Sözleşme Nüshasının Fotoğrafını veya PDF Belgesini Yükleyiniz:", type=["pdf", "png", "jpg", "jpeg"])
        e_notlar = st.text_area("Yaşadığınız Hukuki Sorunu ve Notlarınızı Kısaca Belirtiniz:")

    if st.button("🔒 EVRAKLARI VE ANALİZ TALEBİNİ AVUKATA GÖNDER", use_container_width=True):
        if not e_kiralayan or not e_kiraci:
            st.error("❌ Hata: Sürecin incelenebilmesi için tarafların isim alanları boş bırakılamaz.")
        else:
            dosya_adi = yuklenen_dosya.name if yuklenen_dosya else "Dosya Yüklenmedi"
            conn = sqlite3.connect("hukuk_otomasyon.db")
            cursor = conn.cursor()
            cursor.execute("""
            cursor.execute("""
        INSERT INTO eski_kontratlar 
        (islem_tarihi, kiralayan, kiraci, baslangic_tarihi, aylik_bedel, para_birimi, dosya_adi, notlar)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (islem_tarihi, kiralayan, kiraci, baslangic_tarihi, aylik_bedel, para_birimi, dosya_adi, notlar))
            st.success("✔️ Sözleşme evrakınız ve talebiniz incelenmek üzere Av. Egemen Keskin'in dijital arşivine başarıyla iletilmiştir.")

# =====================================================================
# SENARYO 3: TAMAMEN SANA ÖZEL ŞİFRELİ AVUKAT YÖNETİM PANELİ
# =====================================================================
else:
    st.title("🔒 Keskin Hukuk - Yönetici Otomasyon Kontrol Paneli")
    st.divider()
    
    conn = sqlite3.connect("hukuk_otomasyon.db")
    cursor = conn.cursor()
    cursor.execute("SELECT password_hash FROM admin_auth")
    auth_row = cursor.fetchone()
    conn.close()
    
    # ADIM A: Veri tabanında henüz şifre yoksa (İlk Kurulum ve Şifre Belirleme Ekranı)
    if auth_row is None:
        st.warning("🔑 Egemen Bey, sisteminize ilk kez giriş yapmaktasınız. Lütfen kendinize özel güvenli bir panel şifresi belirleyin:")
        yeni_sifre_1 = st.text_input("Kendi Şifrenizi Oluşturun:", type="password", help="Bu şifre sisteme giriş için anahtarınız olacaktır.")
        yeni_sifre_2 = st.text_input("Şifrenizi Tekrar Yazın:", type="password")
        
        if st.button("🔐 Şifremi Belirle ve Sistemi Kilitle"):
            if yeni_sifre_1 == yeni_sifre_2 and yeni_sifre_1 != "":
                conn = sqlite3.connect("hukuk_otomasyon.db")
                cursor = conn.cursor()
                cursor.execute("INSERT INTO admin_auth (id, password_hash) VALUES (1, ?)", (hash_sifre(yeni_sifre_1),))
                conn.commit()
                conn.close()
                st.success("✔️ Şifreniz başarıyla veri tabanına kilitlendi! Sayfa yenileniyor, artık bu şifreyle giriş yapabilirsiniz.")
                st.rerun()
            else:
                st.error("❌ Hata: Şifreler birbiriyle uyuşmuyor veya boş bırakılamaz!")
                
    # ADIM B: Şifre belirlenmişse (Normal Şifreli Giriş Ekranı)
    else:
        girilen_sifre = st.text_input("Yönetici Giriş Şifrenizi Giriniz:", type="password")
        
        if hash_sifre(girilen_sifre) == auth_row[0]:
            st.success("🔐 Erişim Yetkisi Onaylandı. Hoş geldiniz Avukat Egemen Keskin.")
            
            yonetim_modu = st.radio("Yapmak İstediğiniz İşlemi Seçin:", ["📬 Müvekkillerden Gelen Form Girişleri", "📚 Ofis İçi Eski Kontrat Arşivleme & Analiz"])
            st.divider()
            
            # 1. Müvekkillerden Gelen Form Takibi
            if yonetim_modu == "📬 Müvekkillerden Gelen Form Girişleri":
                st.subheader("📬 Müvekkillerinizin Doldurduğu Yeni Kontrat Talepleri")
                conn = sqlite3.connect("hukuk_otomasyon.db")
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM yeni_kontratlar ORDER BY id DESC")
                yeni_list = cursor.fetchall()
                conn.close()
                
                if not yeni_list:
                    st.info("Sisteme düşen yeni bir müvekkil formu bulunmuyor.")
                else:
                    for k in yeni_list:
                        pb = k['para_birimi'] if 'para_birimi' in k.keys() else "TL (₺)"
                        with st.expander(f"📋 Form #{k['id']} | Kiracı: {k['kiraci']} ({k['tarih']})"):
                            st.write(f"**Kiraya Veren:** {k['kiralayan']} | **T.C:** {k['kiralayan_tc']} | **IBAN:** {k['kiralayan_iban']}")
                            st.write(f"**Kira Şartları:** {k['bedel']:,} {pb} | Artış: %{k['artis']} | Başlangıç: {k['baslangic']}")
                            
                            def word_uret(v, pb_cinsi):
                                doc = docx.Document()
                                p = doc.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p.add_run("KİRA SÖZLEŞMESİ").font.bold = True
                                doc.add_paragraph(f"KİRAYA VEREN: {v['kiralayan']}\nTC: {v['kiralayan_tc']}\nKİRACI: {v['kiraci']}\nTC: {v['kiraci_tc']}\nAdres: {v['kiraci_adres']}")
                                doc.add_paragraph(f"ŞARTLAR: Aylık kira bedeli net {v['bedel']:,} {pb_cinsi}'dir. Her ayın {v['odeme_gunu']}. günü ödenecektir.")
                                bio = io.BytesIO()
                                doc.save(bio)
                                bio.seek(0)
                                return bio
                            
                            w_data = word_uret(k, pb)
                            st.download_button(
                                label="📜 Bu Formu Resmi Word Sözleşmesine Dönüştür ve İndir",
                                data=w_data,
                                file_name=f"Resmi_Kontrat_{k['kiraci']}.docx",
                                key=f"y_btn_{k['id']}"
                            )
            
            # 2. Ofis İçi Eski Kontrat Arşivleme ve Risk Analizi
            else:
                st.subheader("📚 Ofis İçi Eski Kontrat Arşivleme ve Risk Analizi")
                st.info("Bu alan sadece sizin görebileceğiniz, müvekkillere kapalı olan arşiv alanıdır.")
                
                conn = sqlite3.connect("hukuk_otomasyon.db")
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM eski_kontratlar ORDER BY id DESC")
                eski_list = cursor.fetchall()
                conn.close()
                
                if not eski_list:
                    st.info("Arşivde kayıtlı eski sözleşme bulunmuyor. Müvekkilleriniz sol menüden yükleme yaptıkça burası dolacaktır.")
                else:
                    for e in eski_list:
                        epb = e['para_birimi'] if 'para_birimi' in e.keys() else "TL (₺)"
                        with st.expander(f"📁 Arşiv #{e['id']} | {e['kiralayan']} & {e['kiraci']} Kontratı"):
                            st.write(f"**Sözleşme Başlangıç Tarihi:** {e['baslangic_tarihi']}")
                            st.write(f"**Mevcut Aylık Bedel:** {e['aylik_bedel']:,} {epb}")
                            st.write(f"**Ekli Evrak:** 📄 `{e['dosya_adi']}`")
                            st.info(f"📝 **Müvekkilin Notu/Hukuki Sorun:** {e['notlar']}")
                            
                            # Otomatik Borçlar Kanunu Analizi
                            try:
                                bugun = datetime.now().date()
                                baslangic_dt = datetime.strptime(e['baslangic_tarihi'], '%d.%m.%Y').date()
                                fark = bugun - baslangic_dt
                                yil_sayisi = fark.days // 365
                                st.markdown(f"**📊 Hukuki Durum Analizi (TBK uyarınca):**")
                                st.write(f"Sözleşme Yaşı: {yil_sayisi} Yıl")
                                if yil_sayisi >= 5:
                                    st.warning("⚠️ **[TBK m.344/3 - Kira Tespit Davası]:** 5 yıl dolmuş! Rayiç bedel tespiti isteyebilirsiniz.")
                                if yil_sayisi >= 10:
                                    st.error("🚨 **[TBK m.347]:** 10 yıllık uzama süresi dolmuş! Gerekçesiz tahliye ihtar hakkı mevcut.")
                            except:
                                pass
                                
        elif girilen_sifre != "":
            st.error("❌ Geçersiz Yönetici Şifresi! Erişim engellendi.")
